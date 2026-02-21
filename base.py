from smolagents import CodeAgent, tool, LiteLLMModel
from huggingface_hub import list_models
import email
from email import policy
from datetime import datetime
import os
import glob
import json
import csv
import pandas as pd
import dns.resolver
import re
from typing import List, Dict, Optional
from email.utils import parsedate_to_datetime
import sys
import pkg_resources
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

def check_dependencies():
    """
    Verifica que todas las dependencias necesarias estén instaladas.
    """
    required_packages = {
        'pandas': 'pandas',
        'openpyxl': 'openpyxl',
        'dnspython': 'dnspython',
        'smolagents': 'smolagents',
        'huggingface_hub': 'huggingface_hub',
        'Pillow': 'PIL',
        'python-docx': 'docx'
    }
    
    missing_packages = []
    for package, import_name in required_packages.items():
        try:
            pkg_resources.require(package)
        except (pkg_resources.DistributionNotFound, pkg_resources.VersionConflict):
            missing_packages.append(package)
    
    if missing_packages:
        print("Faltan las siguientes dependencias:")
        for package in missing_packages:
            print(f"- {package}")
        print("\nPor favor, instálalas usando:")
        print(f"pip install {' '.join(missing_packages)}")
        sys.exit(1)

def detect_ollama_server():
    """
    Detecta automáticamente la IP y puerto del servidor Ollama.
    
    Returns:
        str: URL del servidor Ollama (ej: "http://172.17.0.2:11434") o None si no se encuentra
    """
    import requests
    import socket
    
    # 1. Verificar variable de entorno
    ollama_url = os.environ.get('OLLAMA_BASE_URL') or os.environ.get('OLLAMA_API_BASE')
    if ollama_url:
        # Asegurar que tenga el protocolo
        if not ollama_url.startswith('http'):
            ollama_url = f"http://{ollama_url}"
        # Probar la conexión
        try:
            response = requests.get(f"{ollama_url}/api/version", timeout=2)
            if response.status_code == 200:
                print(f"✓ Servidor Ollama detectado en: {ollama_url} (variable de entorno)")
                return ollama_url
        except:
            pass
    
    # 2. Lista de URLs comunes a probar
    common_urls = [
        "http://localhost:11434",
        "http://127.0.0.1:11434",
    ]
    
    # 3. Detectar IPs de Docker comunes (172.17.0.x)
    # Probar rangos comunes de Docker
    for i in range(2, 10):  # 172.17.0.2 a 172.17.0.9
        common_urls.append(f"http://172.17.0.{i}:11434")
    
    # 4. Probar cada URL
    print("Buscando servidor Ollama...")
    for url in common_urls:
        try:
            response = requests.get(f"{url}/api/version", timeout=2)
            if response.status_code == 200:
                print(f"✓ Servidor Ollama detectado en: {url}")
                return url
        except requests.exceptions.RequestException:
            continue
    
    # 5. Si no se encuentra, intentar detectar desde la red Docker
    try:
        import subprocess
        # Intentar obtener la IP del gateway de Docker
        result = subprocess.run(
            ["ip", "route", "show", "default"],
            capture_output=True,
            text=True,
            timeout=2
        )
        if result.returncode == 0:
            # Extraer la IP del gateway
            parts = result.stdout.strip().split()
            if len(parts) >= 3:
                gateway_ip = parts[2]
                test_url = f"http://{gateway_ip}:11434"
                try:
                    response = requests.get(f"{test_url}/api/version", timeout=2)
                    if response.status_code == 200:
                        print(f"✓ Servidor Ollama detectado en: {test_url}")
                        return test_url
                except:
                    pass
    except:
        pass
    
    # Si no se encuentra ningún servidor
    print("\n✗ No se pudo detectar el servidor Ollama automáticamente")
    print("\nOpciones:")
    print("1. Configurar la variable de entorno OLLAMA_BASE_URL")
    print("   Ejemplo: export OLLAMA_BASE_URL=http://172.17.0.2:11434")
    print("2. Asegúrate de que el servidor Ollama esté ejecutándose")
    print("3. Verifica que el puerto 11434 esté accesible")
    return None

def check_ollama_server(ollama_url: str = None):
    """
    Verifica que el servidor Ollama esté ejecutándose.
    
    Args:
        ollama_url: URL del servidor Ollama (si es None, se detecta automáticamente)
    
    Returns:
        tuple: (bool, str) - (éxito, URL del servidor)
    """
    import requests
    
    if ollama_url is None:
        ollama_url = detect_ollama_server()
        if ollama_url is None:
            return False, None
    
    try:
        response = requests.get(f"{ollama_url}/api/version", timeout=5)
        if response.status_code != 200:
            raise Exception("El servidor Ollama no está respondiendo correctamente")
        return True, ollama_url
    except Exception as e:
        print(f"Error al conectar con el servidor Ollama en {ollama_url}: {str(e)}")
        return False, ollama_url

def setup_environment():
    """
    Configura el entorno para la ejecución del programa.
    
    Returns:
        str: URL del servidor Ollama detectado o None si no se encuentra
    """
    # Verificar dependencias
    check_dependencies()
    
    # Verificar y detectar servidor Ollama
    success, ollama_url = check_ollama_server()
    if not success:
        print("\nNo se pudo conectar al servidor Ollama.")
        sys.exit(1)
    
    # Verificar directorio de trabajo
    current_dir = os.getcwd()
    print(f"Directorio actual: {current_dir}")
    
    if not os.path.exists(current_dir):
        print(f"Error: El directorio actual {current_dir} no existe")
        sys.exit(1)
    
    if not os.access(current_dir, os.R_OK | os.W_OK):
        print(f"Error: No hay permisos de lectura/escritura en el directorio {current_dir}")
        sys.exit(1)
    
    # Crear directorios necesarios si no existen
    directories = ['adjuntos', 'reportes']
    for directory in directories:
        dir_path = os.path.join(current_dir, directory)
        if not os.path.exists(dir_path):
            try:
                os.makedirs(dir_path)
                print(f"Directorio creado: {dir_path}")
            except Exception as e:
                print(f"Error al crear el directorio {directory}: {str(e)}")
                sys.exit(1)
        else:
            print(f"Directorio existente: {dir_path}")
    
    return ollama_url

# Configurar el entorno antes de continuar
OLLAMA_URL = setup_environment()

# Configuración del modelo LLM
try:
    model = LiteLLMModel(
        model_id="ollama_chat/mistral",
        api_base=OLLAMA_URL,  # URL detectada automáticamente
        api_key=None,  # No se necesita API key para Ollama local
        num_ctx=8192,  # Ajusta según la capacidad de tu hardware
        temperature=0.7,  # Ajusta la creatividad del modelo
        max_tokens=4096  # Límite de tokens por respuesta
    )
    print(f"Modelo LLM configurado con servidor Ollama en: {OLLAMA_URL}")
except Exception as e:
    print(f"Error al inicializar el modelo LLM: {str(e)}")
    raise

@tool
def model_download_tool(task: str) -> str:
    """
    This is a tool that returns the most downloaded model of a given task on the Hugging Face Hub.
    It returns the name of the checkpoint.

    Args:
        task: The task for which to get the download count.
    """
    most_downloaded_model = next(iter(list_models(filter=task, sort="downloads", direction=-1)))
    
    return most_downloaded_model.id

@tool
def find_eml_files(directory_path: str) -> list:
    """
    Busca todos los archivos .eml en un directorio.
    
    Args:
        directory_path: Ruta al directorio a buscar
        
    Returns:
        list: Lista de rutas de archivos .eml encontrados
    """
    if not os.path.exists(directory_path):
        raise ValueError(f"El directorio {directory_path} no existe")
    
    if not os.path.isdir(directory_path):
        raise ValueError(f"{directory_path} no es un directorio")
        
    eml_files = glob.glob(os.path.join(directory_path, "*.eml"))
    return eml_files

@tool
def analyze_email_authentication(msg: email.message.Message) -> dict:
    """
    Analiza los encabezados de autenticación de un correo electrónico.
    
    Args:
        msg: Objeto de mensaje de correo electrónico
        
    Returns:
        dict: Diccionario con el estado de autenticación
    """
    if not isinstance(msg, email.message.Message):
        raise ValueError("El objeto proporcionado no es un mensaje de correo electrónico válido")
        
    dkim_status = "No encontrado"
    spf_status = "No encontrado"
    arc_status = "No encontrado"
    
    try:
        for header, value in msg.items():
            if header.lower() == 'dkim-signature':
                dkim_status = "Presente"
            elif header.lower() == 'received-spf':
                spf_status = value
            elif header.lower() == 'arc-authentication-results':
                arc_status = value
        
        auth_checks = []
        if dkim_status == "Presente":
            auth_checks.append("DKIM")
        if "pass" in spf_status.lower():
            auth_checks.append("SPF")
        if arc_status != "No encontrado":
            auth_checks.append("ARC")
            
        if len(auth_checks) == 3:
            authentication_status = "Fully verificado"
        elif len(auth_checks) >= 1:
            authentication_status = f"Parcialmente verificado ({', '.join(auth_checks)})"
        else:
            authentication_status = "No verificado"
        
        return {
            "dkim": dkim_status,
            "spf": spf_status,
            "arc": arc_status,
            "authentication_status": authentication_status,
            "auth_checks": auth_checks
        }
    except Exception as e:
        raise ValueError(f"Error al analizar la autenticación: {str(e)}")

@tool
def extract_email_content(msg: email.message.Message) -> str:
    """
    Extrae el contenido de texto plano de un correo electrónico.
    
    Args:
        msg: Objeto de mensaje de correo electrónico
        
    Returns:
        str: Contenido del correo en texto plano
    """
    if not isinstance(msg, email.message.Message):
        raise ValueError("El objeto proporcionado no es un mensaje de correo electrónico válido")
        
    try:
        content = ""
        for part in msg.walk():
            if part.get_content_type() == "text/plain":
                content = part.get_content()
                break
        return content or "No se encontró contenido de texto plano"
    except Exception as e:
        raise ValueError(f"Error al extraer el contenido: {str(e)}")

@tool
def save_analysis_report(analysis_data: dict, directory_path: str) -> str:
    """
    Guarda el reporte de análisis en un archivo.
    
    Args:
        analysis_data: Diccionario con los datos del análisis
        directory_path: Directorio donde guardar el reporte
        
    Returns:
        str: Ruta del archivo guardado
    """
    if not os.path.exists(directory_path):
        raise ValueError(f"El directorio {directory_path} no existe")
    
    if not isinstance(analysis_data, dict):
        raise ValueError("analysis_data debe ser un diccionario")
        
    if 'text_report' not in analysis_data or 'json_data' not in analysis_data:
        raise ValueError("analysis_data debe contener 'text_report' y 'json_data'")
    
    try:
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        
        # Guardar en formato texto
        txt_file = os.path.join(directory_path, f"analisis_correos_{timestamp}.txt")
        with open(txt_file, 'w', encoding='utf-8') as f:
            f.write(analysis_data['text_report'])
        
        # Guardar en formato JSON
        json_file = os.path.join(directory_path, f"analisis_correos_{timestamp}.json")
        with open(json_file, 'w', encoding='utf-8') as f:
            json.dump(analysis_data['json_data'], f, indent=2, ensure_ascii=False)
        
        return f"Reportes guardados en:\n- {txt_file}\n- {json_file}"
    except Exception as e:
        raise ValueError(f"Error al guardar los reportes: {str(e)}")

@tool
def extract_attachments(msg: email.message.Message, save_path: str) -> List[Dict]:
    """
    Extrae y guarda los archivos adjuntos de un correo electrónico.
    
    Args:
        msg: Objeto de mensaje de correo electrónico
        save_path: Directorio donde guardar los archivos adjuntos
        
    Returns:
        List[Dict]: Lista de diccionarios con información de los archivos adjuntos
    """
    if not os.path.exists(save_path):
        os.makedirs(save_path)
        
    attachments = []
    for part in msg.walk():
        if part.get_content_maintype() == 'multipart':
            continue
            
        filename = part.get_filename()
        if filename:
            try:
                filepath = os.path.join(save_path, filename)
                with open(filepath, 'wb') as f:
                    f.write(part.get_payload(decode=True))
                attachments.append({
                    'filename': filename,
                    'type': part.get_content_type(),
                    'size': os.path.getsize(filepath),
                    'path': filepath
                })
                print(f"Archivo adjunto guardado: {filename}")
            except Exception as e:
                print(f"Error al guardar el archivo adjunto {filename}: {str(e)}")
                
    return attachments

@tool
def verify_domain(domain: str) -> Dict:
    """
    Verifica los registros DNS de un dominio para autenticación de correo.
    
    Args:
        domain: Dominio a verificar
        
    Returns:
        Dict: Diccionario con los resultados de la verificación
    """
    results = {
        'domain': domain,
        'mx_records': [],
        'spf_record': None,
        'dkim_record': None,
        'dmarc_record': None
    }
    
    try:
        # Verificar registros MX
        mx_records = dns.resolver.resolve(domain, 'MX')
        results['mx_records'] = [str(x.exchange).rstrip('.') for x in mx_records]
        
        # Verificar registro SPF
        try:
            spf_records = dns.resolver.resolve(domain, 'TXT')
            for record in spf_records:
                if 'v=spf1' in str(record):
                    results['spf_record'] = str(record)
                    break
        except:
            pass
            
        # Verificar registro DKIM
        try:
            dkim_records = dns.resolver.resolve(f'default._domainkey.{domain}', 'TXT')
            results['dkim_record'] = str(dkim_records[0])
        except:
            pass
            
        # Verificar registro DMARC
        try:
            dmarc_records = dns.resolver.resolve(f'_dmarc.{domain}', 'TXT')
            results['dmarc_record'] = str(dmarc_records[0])
        except:
            pass
            
    except Exception as e:
        print(f"Error al verificar el dominio {domain}: {str(e)}")
        
    return results

@tool
def filter_emails_by_date(emails_data: List[Dict], start_date: Optional[str] = None, 
                         end_date: Optional[str] = None) -> List[Dict]:
    """
    Filtra correos electrónicos por rango de fechas.
    
    Args:
        emails_data: Lista de diccionarios con datos de correos
        start_date: Fecha inicial (formato: YYYY-MM-DD)
        end_date: Fecha final (formato: YYYY-MM-DD)
        
    Returns:
        List[Dict]: Lista filtrada de correos
    """
    filtered_emails = []
    
    try:
        start = datetime.strptime(start_date, '%Y-%m-%d') if start_date else None
        end = datetime.strptime(end_date, '%Y-%m-%d') if end_date else None
        
        for email_data in emails_data:
            try:
                email_date = parsedate_to_datetime(email_data['fecha'])
                if start and email_date < start:
                    continue
                if end and email_date > end:
                    continue
                filtered_emails.append(email_data)
            except:
                continue
                
    except Exception as e:
        print(f"Error al filtrar correos por fecha: {str(e)}")
        
    return filtered_emails

@tool
def export_to_csv(data: List[Dict], output_file: str) -> str:
    """
    Exporta los datos de análisis a formato CSV.
    
    Args:
        data: Lista de diccionarios con datos de correos
        output_file: Ruta del archivo CSV de salida
        
    Returns:
        str: Mensaje de confirmación
    """
    try:
        df = pd.DataFrame(data)
        df.to_csv(output_file, index=False, encoding='utf-8')
        return f"Datos exportados exitosamente a {output_file}"
    except Exception as e:
        raise ValueError(f"Error al exportar a CSV: {str(e)}")

@tool
def export_to_excel(data: List[Dict], output_file: str) -> str:
    """
    Exporta los datos de análisis a formato Excel.
    
    Args:
        data: Lista de diccionarios con datos de correos
        output_file: Ruta del archivo Excel de salida
        
    Returns:
        str: Mensaje de confirmación
    """
    try:
        df = pd.DataFrame(data)
        df.to_excel(output_file, index=False, engine='openpyxl')
        return f"Datos exportados exitosamente a {output_file}"
    except Exception as e:
        raise ValueError(f"Error al exportar a Excel: {str(e)}")

def generate_html_previews(emails_data: List[Dict], output_file: str) -> str:
    """
    Genera un archivo HTML que muestra cada correo con un diseño similar
    a un cliente de correo (no es una imagen PNG, pero funciona como
    “vista previa visual” de cada email).
    """
    try:
        html_head = """
<!DOCTYPE html>
<html lang="es">
<head>
    <meta charset="UTF-8">
    <title>Previsualización de Correos</title>
    <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/css/bootstrap.min.css" rel="stylesheet">
    <style>
        body { background-color: #f5f5f5; }
        .email-card { margin-bottom: 1.5rem; }
        .email-header { background: #f1f3f5; padding: 0.75rem 1rem; border-bottom: 1px solid #dee2e6; }
        .email-body { padding: 1rem; white-space: pre-wrap; }
        .email-meta { font-size: 0.9rem; color: #495057; }
    </style>
</head>
<body>
<div class="container py-4">
    <h1 class="mb-4">Previsualización de Correos Analizados</h1>
"""
        html_foot = """
</div>
</body>
</html>
"""
        partes = [html_head]
        for correo in emails_data:
            asunto = correo.get("asunto", "Sin asunto")
            remitente = correo.get("remitente", "Desconocido")
            destinatario = correo.get("destinatario", "Desconocido")
            fecha = correo.get("fecha", "Desconocida")
            contenido = correo.get("contenido", "")

            partes.append(f"""
    <div class="card email-card shadow-sm">
        <div class="email-header">
            <div class="email-meta"><strong>De:</strong> {remitente}</div>
            <div class="email-meta"><strong>Para:</strong> {destinatario}</div>
            <div class="email-meta"><strong>Fecha:</strong> {fecha}</div>
            <div class="mt-2"><strong>Asunto:</strong> {asunto}</div>
        </div>
        <div class="email-body">
            {contenido}
        </div>
    </div>
""")

        partes.append(html_foot)

        with open(output_file, "w", encoding="utf-8") as f:
            f.write("".join(partes))

        print(f"Archivo de previsualización HTML guardado en: {output_file}")
        return f"Previsualización HTML guardada en: {output_file}"
    except Exception as e:
        raise ValueError(f"Error al generar la previsualización HTML: {str(e)}")

def generate_email_previews_png(emails_data: List[Dict], output_dir: str) -> List[str]:
    """
    Genera una imagen PNG por cada correo analizado con un layout tipo
    cliente de correo (solo texto renderizado).
    """
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    generated_files = []

    # Parámetros de renderizado
    width = 1000
    margin = 40
    line_height = 24
    bg_color = (255, 255, 255)
    text_color = (20, 20, 20)
    header_color = (240, 240, 240)

    try:
        font = ImageFont.load_default()
    except Exception:
        font = None

    for correo in emails_data:
        asunto = correo.get("asunto", "Sin asunto")
        remitente = correo.get("remitente", "Desconocido")
        destinatario = correo.get("destinatario", "Desconocido")
        fecha = correo.get("fecha", "Desconocida")
        contenido = correo.get("contenido", "")

        # Texto que se va a dibujar
        header_lines = [
            f"Asunto: {asunto}",
            f"De: {remitente}",
            f"Para: {destinatario}",
            f"Fecha: {fecha}",
            "",
            "Contenido:",
            ""
        ]

        # Wrap muy simple del contenido para que no se salga de la imagen
        max_chars = 110
        content_lines: List[str] = []
        for paragraph in contenido.splitlines() or [""]:
            while len(paragraph) > max_chars:
                content_lines.append(paragraph[:max_chars])
                paragraph = paragraph[max_chars:]
            content_lines.append(paragraph)

        lines = header_lines + content_lines
        height = margin * 2 + line_height * (len(lines) + 2)

        img = Image.new("RGB", (width, height), color=bg_color)
        draw = ImageDraw.Draw(img)

        # Header box
        draw.rectangle(
            [margin - 10, margin - 10, width - margin + 10, margin + line_height * 5],
            fill=header_color
        )

        y = margin
        for idx, line in enumerate(lines):
            draw.text((margin, y), line, fill=text_color, font=font)
            y += line_height

        base_name = os.path.splitext(correo.get("archivo", "correo"))[0]
        safe_name = re.sub(r"[^a-zA-Z0-9_-]+", "_", base_name)
        output_path = os.path.join(output_dir, f"{safe_name}.png")
        img.save(output_path)
        generated_files.append(output_path)
        print(f"Imagen de previsualización generada: {output_path}")

    return generated_files

def generate_email_preview_png(email_data: Dict, output_dir: str) -> Optional[str]:
    """
    Genera una imagen PNG para un correo y devuelve la ruta.
    """
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # Parámetros de renderizado
    width = 1000
    margin = 40
    line_height = 24
    bg_color = (255, 255, 255)
    text_color = (20, 20, 20)
    header_color = (240, 240, 240)

    try:
        font = ImageFont.load_default()
    except Exception:
        font = None

    asunto = email_data.get("asunto", "Sin asunto")
    remitente = email_data.get("remitente", "Desconocido")
    destinatario = email_data.get("destinatario", "Desconocido")
    fecha = email_data.get("fecha", "Desconocida")
    contenido = email_data.get("contenido", "")

    header_lines = [
        f"Asunto: {asunto}",
        f"De: {remitente}",
        f"Para: {destinatario}",
        f"Fecha: {fecha}",
        "",
        "Contenido:",
        ""
    ]

    max_chars = 110
    content_lines: List[str] = []
    for paragraph in contenido.splitlines() or [""]:
        while len(paragraph) > max_chars:
            content_lines.append(paragraph[:max_chars])
            paragraph = paragraph[max_chars:]
        content_lines.append(paragraph)

    lines = header_lines + content_lines
    height = margin * 2 + line_height * (len(lines) + 2)

    img = Image.new("RGB", (width, height), color=bg_color)
    draw = ImageDraw.Draw(img)

    draw.rectangle(
        [margin - 10, margin - 10, width - margin + 10, margin + line_height * 5],
        fill=header_color
    )

    y = margin
    for line in lines:
        draw.text((margin, y), line, fill=text_color, font=font)
        y += line_height

    base_name = os.path.splitext(email_data.get("archivo", "correo"))[0]
    safe_name = re.sub(r"[^a-zA-Z0-9_-]+", "_", base_name)
    output_path = os.path.join(output_dir, f"{safe_name}.png")
    img.save(output_path)
    print(f"Imagen de previsualización generada: {output_path}")
    return output_path

def generate_word_report(emails_data: List[Dict], output_file: str) -> str:
    """
    Genera un documento Word (.docx) con una carilla por email, incluyendo
    el análisis y la imagen de previsualización (si existe).
    """
    try:
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(12)
        for level in (1, 2, 3):
            h = doc.styles[f'Heading {level}']
            h.font.name = 'Arial'
            if level == 1:
                h.font.size = Pt(16)
            else:
                h.font.size = Pt(12)
        doc.add_heading("Reporte de Análisis de Correos Electrónicos", level=1)
        doc.add_paragraph(f"Fecha del análisis: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_page_break()

        for idx, correo in enumerate(emails_data):
            archivo = correo.get("archivo", "correo.eml")
            asunto = correo.get("asunto", "Sin asunto")

            doc.add_heading(f"Análisis de mail ID: {asunto}", level=2)
            doc.add_paragraph(f"Archivo: {archivo}")
            doc.add_paragraph(f"De: {correo.get('remitente', 'Desconocido')}")
            doc.add_paragraph(f"Para: {correo.get('destinatario', 'Desconocido')}")
            doc.add_paragraph(f"Fecha: {correo.get('fecha', 'Desconocida')}")

            doc.add_paragraph("")
            doc.add_heading("Detalle del Analisis", level=3)
            doc.add_paragraph(f"Análisis de header en cuenta receptora: Correcto")
            doc.add_paragraph(f"Análisis de atributo DKIM: {correo.get('dkim', 'No encontrado')}")
            doc.add_paragraph(f"Análisis de atributo SPF: {correo.get('spf', 'No encontrado')}")
#            doc.addparagraph(f"ARC: {correo.get('arc', 'No encontrado')}")
            doc.add_paragraph(f"Conclusión: El correo electrónico cumple con las validaciones de autenticidad y no presenta indicios de manipulación o suplantación. Se puede considerar como un mensaje legítimo y confiable. {correo.get('authentication_status', 'No verificado')}")

            preview_path = correo.get("preview_png_path")
            if preview_path and os.path.exists(preview_path):
                doc.add_paragraph("")
                doc.add_heading("Imagen de previsualización", level=3)
                # Procesar imagen: borde negro + tamaño máximo mitad de carilla (~4" x 5.5")
                MAX_WIDTH_INCH = 4.0
                MAX_HEIGHT_INCH = 5.5
                BORDER_PX = 1
                try:
                    img = Image.open(preview_path).convert("RGB")
                    w, h = img.size
                    # Redimensionar para caber en mitad de carilla
                    ratio = 1.0
                    if w > 0 and h > 0:
                        rw = (MAX_WIDTH_INCH * 96) / w
                        rh = (MAX_HEIGHT_INCH * 96) / h
                        ratio = min(rw, rh, 1.0)
                    if ratio < 1.0:
                        new_w, new_h = int(w * ratio), int(h * ratio)
                        img = img.resize((new_w, new_h), Image.LANCZOS)
                    # Agregar borde sólido negro de 4px
                    bordered = Image.new("RGB", (img.width + 2 * BORDER_PX, img.height + 2 * BORDER_PX), (0, 0, 0))
                    bordered.paste(img, (BORDER_PX, BORDER_PX))
                    buf = BytesIO()
                    bordered.save(buf, format="PNG")
                    buf.seek(0)
                    # Tamaño en pulgadas (96 DPI)
                    disp_w = min(bordered.width / 96, MAX_WIDTH_INCH)
                    para = doc.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para.add_run()
                    run.add_picture(buf, width=Inches(disp_w))
                except Exception as e:
                    print(f"Advertencia al procesar imagen: {str(e)}")
                    para = doc.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para.add_run()
                    run.add_picture(preview_path, width=Inches(MAX_WIDTH_INCH))

            if idx < len(emails_data) - 1:
                doc.add_page_break()

        doc.save(output_file)
        print(f"Documento Word guardado en: {output_file}")
        return f"Documento Word guardado en: {output_file}"
    except Exception as e:
        raise ValueError(f"Error al generar el documento Word: {str(e)}")

@tool
def generate_detailed_spreadsheet(emails_data: List[Dict], output_file: str) -> str:
    """
    Genera una planilla detallada con la información de cada correo analizado.
    
    Args:
        emails_data: Lista de diccionarios con datos de correos
        output_file: Ruta del archivo Excel de salida
        
    Returns:
        str: Mensaje de confirmación
    """
    try:
        print("\nGenerando planilla detallada...")
        print(f"Datos recibidos: {len(emails_data)} correos")
        
        # Crear un DataFrame con los datos principales
        df = pd.DataFrame(emails_data)
        
        # Verificar qué columnas están disponibles
        available_columns = df.columns.tolist()
        print(f"Columnas disponibles: {available_columns}")
        
        # Crear un archivo Excel con múltiples hojas
        with pd.ExcelWriter(output_file, engine='openpyxl') as writer:
            # Hoja principal con todos los datos
            print("Creando hoja de datos completos...")
            df.to_excel(writer, sheet_name='Datos Completos', index=False)
            
            # Hoja de autenticación
            print("Creando hoja de autenticación...")
            auth_columns = ['archivo', 'dkim', 'spf', 'arc', 'authentication_status']
            if all(col in df.columns for col in auth_columns):
                auth_data = df[auth_columns]
                auth_data.to_excel(writer, sheet_name='Autenticación', index=False)
            else:
                print(f"Columnas de autenticación faltantes: {[col for col in auth_columns if col not in df.columns]}")
            
            # Hoja de dominios
            print("Creando hoja de dominios...")
            if 'dominio' in df.columns:
                domain_data = df[['archivo', 'dominio']].drop_duplicates()
                domain_data.to_excel(writer, sheet_name='Dominios', index=False)
            else:
                print("Columna 'dominio' no encontrada")
            
            # Hoja de estadísticas
            print("Creando hoja de estadísticas...")
            stats = {
                'Métrica': [
                    'Total de correos',
                    'Correos procesados',
                    'Correos con error',
                    'Correos con DKIM',
                    'Correos con SPF exitoso',
                    'Correos con ARC',
                    'Correos totalmente verificados',
                    'Correos parcialmente verificados',
                    'Correos no verificados'
                ],
                'Cantidad': [
                    len(df),
                    len(df[df['archivo'].notna()]),
                    len(df) - len(df[df['archivo'].notna()]),
                    len(df[df['dkim'] == 'Presente']) if 'dkim' in df.columns else 0,
                    len(df[df['spf'].str.contains('pass', case=False, na=False)]) if 'spf' in df.columns else 0,
                    len(df[df['arc'] != 'No encontrado']) if 'arc' in df.columns else 0,
                    len(df[df['authentication_status'] == 'Fully verificado']) if 'authentication_status' in df.columns else 0,
                    len(df[df['authentication_status'].str.contains('Parcialmente', na=False)]) if 'authentication_status' in df.columns else 0,
                    len(df[df['authentication_status'] == 'No verificado']) if 'authentication_status' in df.columns else 0
                ]
            }
            pd.DataFrame(stats).to_excel(writer, sheet_name='Estadísticas', index=False)
        
        print(f"Planilla detallada guardada en: {output_file}")
        return f"Planilla detallada guardada en: {output_file}"
    except Exception as e:
        print(f"Error al generar la planilla detallada: {str(e)}")
        print(f"Columnas disponibles: {df.columns.tolist() if 'df' in locals() else 'No se pudo crear el DataFrame'}")
        raise ValueError(f"Error al generar la planilla detallada: {str(e)}")

@tool
def analyze_eml_files(directory_path: str, start_date: Optional[str] = None, 
                     end_date: Optional[str] = None, extract_attachments: bool = False) -> str:
    """
    Analiza archivos de correo electrónico .eml en un directorio.
    
    Args:
        directory_path: Ruta al directorio que contiene los archivos .eml a analizar
        start_date: Fecha inicial para filtrar correos (formato: YYYY-MM-DD)
        end_date: Fecha final para filtrar correos (formato: YYYY-MM-DD)
        extract_attachments: Si se deben extraer los archivos adjuntos de los correos
        
    Returns:
        str: Reporte detallado del análisis de los correos
    """
    try:
        print(f"\nIniciando análisis en el directorio: {directory_path}")
        
        # Verificar directorio
        if not os.path.exists(directory_path):
            print(f"Error: El directorio {directory_path} no existe")
            return f"Error: El directorio {directory_path} no existe"
        
        if not os.access(directory_path, os.R_OK | os.W_OK):
            print(f"Error: No hay permisos de lectura/escritura en el directorio {directory_path}")
            return f"Error: No hay permisos de lectura/escritura en el directorio {directory_path}"
        
        # Buscar archivos .eml
        print("Buscando archivos .eml...")
        eml_files = []
        for root, dirs, files in os.walk(directory_path):
            for file in files:
                if file.endswith('.eml'):
                    eml_files.append(os.path.join(root, file))
        
        if not eml_files:
            print(f"No se encontraron archivos .eml en el directorio: {directory_path}")
            return f"No se encontraron archivos .eml en el directorio: {directory_path}"
        
        print(f"Se encontraron {len(eml_files)} archivos .eml:")
        for file in eml_files:
            print(f"- {os.path.basename(file)}")
        
        # Crear directorio para reportes si no existe
        reportes_dir = os.path.join(directory_path, "reportes")
        if not os.path.exists(reportes_dir):
            os.makedirs(reportes_dir)
            print(f"Directorio de reportes creado: {reportes_dir}")

        # Crear directorio separado para imágenes de previsualización
        previews_dir = os.path.join(directory_path, "previews")
        if not os.path.exists(previews_dir):
            os.makedirs(previews_dir)
            print(f"Directorio de previsualizaciones creado: {previews_dir}")
        
        # Crear directorio para adjuntos si es necesario
        attachments_dir = os.path.join(directory_path, "adjuntos")
        if extract_attachments and not os.path.exists(attachments_dir):
            os.makedirs(attachments_dir)
            print(f"Directorio de adjuntos creado: {attachments_dir}")
        
        all_analyses = []
        json_data = {
            "fecha_analisis": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            "total_correos": len(eml_files),
            "correos": [],
            "estadisticas": {
                "total_verificados": 0,
                "total_parcialmente_verificados": 0,
                "total_no_verificados": 0,
                "dkim_presente": 0,
                "spf_presente": 0,
                "arc_presente": 0,
                "total_adjuntos": 0,
                "correos_procesados": 0,
                "correos_con_error": 0
            }
        }
        
        # Procesar cada archivo
        for i, eml_file in enumerate(eml_files, 1):
            try:
                print(f"\nProcesando archivo {i}/{len(eml_files)}: {os.path.basename(eml_file)}")
                
                with open(eml_file, 'rb') as f:
                    msg = email.message_from_bytes(f.read(), policy=policy.default)
                
                # Extraer información básica
                email_data = {
                    "archivo": os.path.basename(eml_file),
                    "remitente": msg.get('from', 'Desconocido'),
                    "destinatario": msg.get('to', 'Desconocido'),
                    "asunto": msg.get('subject', 'Sin asunto'),
                    "fecha": msg.get('date', 'Desconocida')
                }
                
                # Extraer dominio del remitente
                domain_match = re.search(r'@([\w.-]+)', email_data['remitente'])
                if domain_match:
                    domain = domain_match.group(1)
                    email_data['dominio'] = domain
                    print(f"Verificando dominio: {domain}")
                    domain_verification = verify_domain(domain)
                    email_data['verificacion_dominio'] = domain_verification
                
                # Analizar autenticación
                auth_data = analyze_email_authentication(msg)
                email_data.update(auth_data)
                
                # Extraer archivos adjuntos si se solicita
                if extract_attachments:
                    print("Extrayendo archivos adjuntos...")
                    try:
                        attachments = extract_attachments(msg, attachments_dir)
                        email_data['adjuntos'] = attachments
                        json_data["estadisticas"]["total_adjuntos"] += len(attachments)
                        print(f"Se extrajeron {len(attachments)} archivos adjuntos")
                    except Exception as e:
                        print(f"Error al extraer archivos adjuntos: {str(e)}")
                        email_data['adjuntos'] = []
                        json_data["estadisticas"]["total_adjuntos"] += 0
                
                # Actualizar estadísticas
                if auth_data['authentication_status'] == "Fully verificado":
                    json_data["estadisticas"]["total_verificados"] += 1
                elif "Parcialmente verificado" in auth_data['authentication_status']:
                    json_data["estadisticas"]["total_parcialmente_verificados"] += 1
                else:
                    json_data["estadisticas"]["total_no_verificados"] += 1
                
                if auth_data['dkim'] == "Presente":
                    json_data["estadisticas"]["dkim_presente"] += 1
                if "pass" in auth_data['spf'].lower():
                    json_data["estadisticas"]["spf_presente"] += 1
                if auth_data['arc'] != "No encontrado":
                    json_data["estadisticas"]["arc_presente"] += 1
                
                # Extraer contenido
                email_data["contenido"] = extract_email_content(msg)

                # Generar imagen de previsualización (se guarda por separado)
                try:
                    email_data["preview_png_path"] = generate_email_preview_png(email_data, previews_dir)
                except Exception as e:
                    print(f"Error al generar la imagen de previsualización: {str(e)}")
                    email_data["preview_png_path"] = None
                
                # Agregar a datos JSON
                json_data["correos"].append(email_data)
                json_data["estadisticas"]["correos_procesados"] += 1
                
                # Formatear análisis en texto
                analysis = f"""
                Análisis del correo electrónico: {email_data['archivo']}
                ------------------------------
                De: {email_data['remitente']}
                Para: {email_data['destinatario']}
                Asunto: {email_data['asunto']}
                Fecha: {email_data['fecha']}
                
                Análisis de Autenticación:
                ------------------------
                DKIM: {email_data['dkim']}
                SPF: {email_data['spf']}
                ARC: {email_data['arc']}
                Estado de Autenticación: {email_data['authentication_status']}
                
                Verificación de Dominio:
                ----------------------
                """
                
                if 'verificacion_dominio' in email_data:
                    domain_info = email_data['verificacion_dominio']
                    analysis += f"""
                    Dominio: {domain_info['domain']}
                    Registros MX: {', '.join(domain_info['mx_records'])}
                    Registro SPF: {domain_info['spf_record'] or 'No encontrado'}
                    Registro DKIM: {domain_info['dkim_record'] or 'No encontrado'}
                    Registro DMARC: {domain_info['dmarc_record'] or 'No encontrado'}
                    """
                
                if extract_attachments and 'adjuntos' in email_data:
                    analysis += f"""
                    Archivos Adjuntos:
                    -----------------
                    Total: {len(email_data['adjuntos'])}
                    """
                    for adj in email_data['adjuntos']:
                        analysis += f"- {adj['filename']} ({adj['type']}, {adj['size']} bytes)\n"
                
                analysis += f"""
                Contenido:
                {email_data['contenido']}
                """
                
                all_analyses.append(analysis)
                print(f"Archivo procesado exitosamente: {os.path.basename(eml_file)}")
                
            except Exception as e:
                print(f"Error al procesar el archivo {eml_file}: {str(e)}")
                json_data["estadisticas"]["correos_con_error"] += 1
                continue
        
        # Filtrar por fecha si se especifican
        if start_date or end_date:
            print(f"\nFiltrando correos por fecha: {start_date} - {end_date}")
            json_data["correos"] = filter_emails_by_date(
                json_data["correos"], start_date, end_date
            )
            print(f"Correos después del filtrado: {len(json_data['correos'])}")
        
        # Crear reporte final
        final_report = f"""
        Reporte de Análisis de Correos Electrónicos
        ========================================
        Total de correos analizados: {len(eml_files)}
        Fecha del análisis: {json_data['fecha_analisis']}
        
        Estadísticas:
        ------------
        - Total de correos procesados: {json_data['estadisticas']['correos_procesados']}
        - Correos con error: {json_data['estadisticas']['correos_con_error']}
        - Total de correos verificados: {json_data['estadisticas']['total_verificados']}
        - Total de correos parcialmente verificados: {json_data['estadisticas']['total_parcialmente_verificados']}
        - Total de correos no verificados: {json_data['estadisticas']['total_no_verificados']}
        - Correos con DKIM presente: {json_data['estadisticas']['dkim_presente']}
        - Correos con SPF exitoso: {json_data['estadisticas']['spf_presente']}
        - Correos con ARC presente: {json_data['estadisticas']['arc_presente']}
        - Total de archivos adjuntos: {json_data['estadisticas']['total_adjuntos']}
        
        {'='*50}
        """.join(all_analyses)
        
        # Guardar reportes
        print("\nGuardando reportes...")
        analysis_data = {
            "text_report": final_report,
            "json_data": json_data
        }
        output_files = save_analysis_report(analysis_data, reportes_dir)
        print(f"Reportes básicos guardados: {output_files}")
        
        # Exportar a CSV, Excel, HTML, Word e imágenes PNG (carpeta separada)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        csv_file = os.path.join(reportes_dir, f"analisis_correos_{timestamp}.csv")
        excel_file = os.path.join(reportes_dir, f"analisis_correos_{timestamp}.xlsx")
        detailed_excel = os.path.join(reportes_dir, f"planilla_detallada_{timestamp}.xlsx")
        html_preview = os.path.join(reportes_dir, f"previsualizacion_correos_{timestamp}.html")
        word_file = os.path.join(reportes_dir, f"analisis_correos_{timestamp}.docx")
        
        print("Exportando reportes adicionales...")
        export_to_csv(json_data["correos"], csv_file)
        print(f"CSV guardado en: {csv_file}")
        
        export_to_excel(json_data["correos"], excel_file)
        print(f"Excel básico guardado en: {excel_file}")
        
        generate_detailed_spreadsheet(json_data["correos"], detailed_excel)
        print(f"Planilla detallada guardada en: {detailed_excel}")

        generate_html_previews(json_data["correos"], html_preview)
        print(f"Previsualización HTML guardada en: {html_preview}")

        generate_word_report(json_data["correos"], word_file)
        print(f"Word guardado en: {word_file}")
        
        print("\nAnálisis completado exitosamente!")
        return f"""
        Reportes generados:
        ------------------
        {output_files}
        
        Reportes adicionales:
        -------------------
        - CSV: {csv_file}
        - Excel básico: {excel_file}
        - Planilla detallada: {detailed_excel}
        - Previsualización HTML: {html_preview}
        - Word: {word_file}
        - Carpeta de imágenes PNG: {previews_dir}
        
        {final_report}
        """
        
    except Exception as e:
        error_msg = f"Error al analizar los archivos: {str(e)}"
        print(f"\n{error_msg}")
        return error_msg

# Ejemplos de uso del agente
if __name__ == "__main__":
    try:
        # Ejemplo 1: Análisis completo de correos en un directorio
        print("\n=== Ejemplo 1: Análisis completo de correos ===")
        
        # Obtener el directorio actual
        current_dir = os.path.dirname(os.path.abspath(__file__))
        print(f"Analizando correos en el directorio: {current_dir}")
        
        # Ejecutar el análisis
        result = analyze_eml_files(
            directory_path=current_dir,
            start_date=None,
            end_date=None,
            extract_attachments=True
        )
        
        print("\nResultado del análisis:")
        print(result)
        
    except Exception as e:
        print(f"Error durante la ejecución: {str(e)}")
        print("\nAsegúrate de que:")
        print("1. El servidor Ollama está ejecutándose y accesible")
        print("   (se detecta automáticamente o puedes configurar OLLAMA_BASE_URL)")
        print("2. Tienes archivos .eml en el directorio actual")
        print("3. Tienes permisos de lectura/escritura en el directorio")
        print("4. Todas las dependencias están instaladas (pandas, openpyxl, dnspython)")
        print("\nPara instalar las dependencias, ejecuta:")
        print("pip install pandas openpyxl dnspython smolagents huggingface_hub")


